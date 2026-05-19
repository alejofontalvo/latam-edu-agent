import hashlib
import shutil
from pathlib import Path
from typing import Any
from datetime import datetime

import chromadb
from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models.entities import Document


settings = get_settings()


class HashEmbeddingFunction:
    """Embedding local determinístico para una versión académica sin API externa."""

    def __call__(self, input: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in input]

    @staticmethod
    def _embed(text: str, dimensions: int = 384) -> list[float]:
        vector = [0.0] * dimensions
        for token in text.lower().split():
            digest = hashlib.sha256(token.encode("utf-8")).digest()
            index = int.from_bytes(digest[:4], "little") % dimensions
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vector[index] += sign
        norm = sum(value * value for value in vector) ** 0.5 or 1.0
        return [value / norm for value in vector]


def get_collection():
    client = chromadb.PersistentClient(path=str(settings.chroma_dir))
    return client.get_or_create_collection(
        name=settings.rag_collection_name,
        embedding_function=HashEmbeddingFunction(),
    )


def extract_text(path: str | Path) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    raise ValueError("Formato documental no soportado. Usa PDF, TXT, Markdown o DOCX.")


def chunk_documents(text: str, chunk_size: int | None = None, overlap: int | None = None) -> list[str]:
    size = chunk_size or settings.chunk_size
    step_back = overlap or settings.chunk_overlap
    clean = " ".join(text.split())
    if not clean:
        return []

    chunks = []
    start = 0
    while start < len(clean):
        end = start + size
        chunks.append(clean[start:end])
        start = end - step_back if end < len(clean) else end
    return chunks


async def upload_documents(files: list[UploadFile], db: Session) -> dict:
    collection = get_collection()
    uploaded = []
    total_chunks = 0

    for file in files:
        filename = Path(file.filename or "document.txt").name.replace(" ", "_")
        document_path = settings.documents_dir / filename
        with document_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        text = extract_text(document_path)
        chunks = chunk_documents(text)
        if not chunks:
            continue

        ids = [f"{filename}-{index}" for index in range(len(chunks))]
        metadatas = [
            {"source": filename, "chunk": index, "path": str(document_path)}
            for index in range(len(chunks))
        ]
        collection.upsert(ids=ids, documents=chunks, metadatas=metadatas)

        db.add(
            Document(
                filename=filename,
                title=document_path.stem.replace("_", " ").title(),
                source="Carga administrativa",
                path=str(document_path),
                chunk_count=len(chunks),
            )
        )
        uploaded.append({"filename": filename, "chunks": len(chunks)})
        total_chunks += len(chunks)

    db.commit()
    return {"uploaded": uploaded, "total_chunks": total_chunks}


def index_documents_directory(db: Session) -> dict:
    collection = get_collection()
    supported = {".pdf", ".txt", ".md", ".docx"}
    indexed = []
    total_chunks = 0
    for document_path in settings.documents_dir.rglob("*"):
        if not document_path.is_file() or document_path.suffix.lower() not in supported:
            continue
        text = extract_text(document_path)
        chunks = chunk_documents(text)
        if not chunks:
            continue
        source_group = document_path.parent.name
        filename = document_path.name
        ids = [f"{source_group}-{filename}-{index}" for index in range(len(chunks))]
        metadatas = [
            {
                "source": filename,
                "title": document_path.stem.replace("_", " ").title(),
                "chunk": index,
                "path": str(document_path),
                "source_group": source_group,
            }
            for index in range(len(chunks))
        ]
        collection.upsert(ids=ids, documents=chunks, metadatas=metadatas)
        existing = db.query(Document).filter(Document.path == str(document_path)).first()
        document = existing or Document(filename=filename, path=str(document_path))
        document.title = document_path.stem.replace("_", " ").title()
        document.source = source_group
        document.document_type = "metodologico"
        document.chunk_count = len(chunks)
        document.created_at = datetime.utcnow()
        if not existing:
            db.add(document)
        indexed.append({"filename": filename, "source": source_group, "chunks": len(chunks)})
        total_chunks += len(chunks)
    db.commit()
    return {"indexed": indexed, "total_chunks": total_chunks}


def list_documents(db: Session) -> list[dict]:
    return [
        {
            "id": item.id,
            "filename": item.filename,
            "title": item.title or item.filename,
            "source": item.source,
            "country": item.country,
            "exam": item.exam,
            "year": item.year,
            "document_type": item.document_type,
            "chunk_count": item.chunk_count,
            "created_at": item.created_at.isoformat(),
        }
        for item in db.query(Document).order_by(Document.created_at.desc()).all()
    ]


def retrieve_context(question: str, top_k: int = 5) -> dict[str, Any]:
    collection = get_collection()
    results = collection.query(query_texts=[question], n_results=top_k)
    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]
    distances = results.get("distances", [[]])[0]

    sources = []
    context_parts = []
    for index, document in enumerate(documents):
        metadata = metadatas[index] if index < len(metadatas) else {}
        distance = distances[index] if index < len(distances) else None
        context_parts.append(f"Fuente {index + 1} ({metadata.get('source')}): {document}")
        sources.append(
            {
                "source": metadata.get("source"),
                "chunk": metadata.get("chunk"),
                "distance": distance,
                "excerpt": document[:280],
            }
        )
    return {"context": "\n\n".join(context_parts), "sources": sources}


def generate_answer_with_context(question: str, context: str) -> str:
    if not context:
        return (
            "No encontré documentos indexados que respalden la respuesta. "
            "Carga informes, diccionarios de variables o guías metodológicas para usar RAG."
        )
    return (
        "Con base en los documentos recuperados, la respuesta debe interpretarse así:\n\n"
        f"{context[:1800]}\n\n"
        "Síntesis: revisa las fuentes citadas para validar definiciones, metodología y alcance. "
        "No se reportan cifras nuevas si no provienen del dataset o del documento cargado."
    )
